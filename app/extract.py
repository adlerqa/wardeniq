"""Pull plain text out of uploaded requirement docs: PDF, DOCX, Markdown/text;
and source files out of a GitHub repo tarball."""
import html as _htmllib
import io
import os
import re
import tarfile
from html.parser import HTMLParser


class _TextExtractor(HTMLParser):
    """Very small HTML→text pass for Confluence storage-format bodies."""
    _BLOCK = {"p", "br", "li", "tr", "div", "h1", "h2", "h3", "h4", "h5", "h6"}

    def __init__(self):
        super().__init__()
        self._parts = []
        self._skip = 0

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style"):
            self._skip += 1
        elif tag in self._BLOCK:
            self._parts.append("\n")

    def handle_endtag(self, tag):
        if tag in ("script", "style") and self._skip:
            self._skip -= 1

    def handle_data(self, data):
        if not self._skip and data.strip():
            self._parts.append(data)


def html_to_text(html_str: str) -> str:
    """Convert HTML (e.g. Confluence storage format) to readable plain text."""
    if not html_str:
        return ""
    p = _TextExtractor()
    try:
        p.feed(html_str)
        raw = "".join(p._parts)
    except Exception:  # noqa: BLE001  -- malformed markup: strip tags crudely
        raw = re.sub(r"<[^>]+>", " ", html_str)
    text = _htmllib.unescape(raw)
    lines = [re.sub(r"[ \t]+", " ", ln).strip() for ln in text.splitlines()]
    return "\n".join(ln for ln in lines if ln).strip()


_URL_RE = re.compile(r"https?://[^\s)>\]\"'}]+")


def urls_in_text(text: str) -> list:
    """All http(s) URLs written as visible text."""
    return _URL_RE.findall(text or "")


def pdf_links(data: bytes) -> list:
    """Every http(s) URL in a PDF — both clickable link annotations and visible-text URLs."""
    urls = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        for page in reader.pages:
            annots = page.get("/Annots")
            if annots:
                try:
                    for a in annots:
                        obj = a.get_object()
                        act = obj.get("/A")
                        if act is not None:
                            act = act.get_object() if hasattr(act, "get_object") else act
                            uri = act.get("/URI") if hasattr(act, "get") else None
                            if uri:
                                urls.append(str(uri))
                except Exception:  # noqa: BLE001
                    pass
            try:
                urls.extend(urls_in_text(page.extract_text() or ""))
            except Exception:  # noqa: BLE001
                pass
    except Exception:  # noqa: BLE001
        pass
    # de-dupe, keep order, http(s) only
    out, seen = [], set()
    for u in urls:
        u = u.strip().rstrip(".,);")
        if u.startswith("http") and u not in seen:
            seen.add(u)
            out.append(u)
    return out


class _LinkExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.links = []

    def handle_starttag(self, tag, attrs):
        if tag == "a":
            for k, v in attrs:
                if k == "href" and v:
                    self.links.append(v)


def extract_html_links(html_str: str, base_url: str = "") -> list:
    """Absolute http(s) links found in an HTML page (relative links resolved against base_url)."""
    from urllib.parse import urljoin, urlparse
    p = _LinkExtractor()
    try:
        p.feed(html_str or "")
    except Exception:  # noqa: BLE001
        return []
    out, seen = [], set()
    for href in p.links:
        u = urljoin(base_url, href) if base_url else href
        if urlparse(u).scheme in ("http", "https") and u not in seen:
            seen.add(u)
            out.append(u)
    return out


def confluence_page_id_from_url(url: str):
    """Extract the numeric page id from a Confluence URL (…/pages/<id>/… or ?pageId=<id>),
    or accept a bare numeric id."""
    if not url:
        return None
    m = re.search(r"/pages/(\d+)", url) or re.search(r"[?&]pageId=(\d+)", url)
    if m:
        return m.group(1)
    u = url.strip()
    return u if u.isdigit() else None

CODE_EXT = {".py", ".js", ".ts", ".tsx", ".jsx", ".mjs", ".cjs", ".cts", ".mts",
            ".go", ".java", ".rb", ".cs", ".php", ".rs", ".kt", ".kts", ".swift",
            ".c", ".cc", ".cpp", ".cxx", ".h", ".hpp", ".scala", ".m", ".mm",
            ".sql", ".vue", ".svelte", ".astro", ".sh", ".bash", ".dart", ".ex",
            ".exs", ".clj", ".cljs", ".groovy", ".gradle", ".lua", ".pl", ".pm",
            ".r", ".jl", ".erl", ".hs", ".fs", ".vb"}
SKIP_DIRS = {"node_modules", "vendor", "dist", "build", ".git", "__pycache__", ".next",
             ".venv", "venv", "target", ".idea", "coverage", "out", ".cache",
             "bin", "obj", ".nuxt", ".svelte-kit"}


def source_files_from_tar(data: bytes, max_files=400, max_bytes=40000, return_stats=False):
    """Return [(relative_path, text)] of source files from a GitHub tarball.

    With return_stats=True, also returns a diagnostics dict so callers can explain an
    empty result (total files seen, extension histogram, sample paths)."""
    out = []
    total_files = 0
    ext_counts = {}
    sample = []
    with tarfile.open(fileobj=io.BytesIO(data), mode="r:gz") as tf:
        for m in tf.getmembers():
            if not m.isfile():
                continue
            total_files += 1
            parts = m.name.split("/", 1)          # strip the <owner>-<repo>-<sha>/ prefix
            rel = parts[1] if len(parts) > 1 else m.name
            if len(sample) < 10 and rel:
                sample.append(rel)
            if not rel or any(seg in SKIP_DIRS for seg in rel.split("/")):
                continue
            ext = os.path.splitext(rel)[1].lower() or "(none)"
            ext_counts[ext] = ext_counts.get(ext, 0) + 1
            if ext not in CODE_EXT:
                continue
            if m.size > max_bytes * 6:            # skip huge generated files
                continue
            f = tf.extractfile(m)
            if not f:
                continue
            try:
                txt = f.read().decode("utf-8", "replace")
            except Exception:  # noqa: BLE001
                continue
            out.append((rel, txt[:max_bytes]))
            if len(out) >= max_files:
                break
    if return_stats:
        top_ext = sorted(ext_counts.items(), key=lambda x: -x[1])[:12]
        return out, {"total_files": total_files, "top_ext": top_ext, "sample": sample}
    return out


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _pdf(data)
    if name.endswith(".docx"):
        return _docx(data)
    # md / txt / anything else -> decode as text
    return data.decode("utf-8", errors="replace")


def _pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            pass
    return "\n".join(parts).strip()


def _docx(data: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    # include table cells too
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def chunk(text: str, max_chars: int = 1200, overlap: int = 150) -> list[str]:
    """Coarse chunking for embedding the source doc (used for feature-level RAG)."""
    text = text.strip()
    if len(text) <= max_chars:
        return [text] if text else []
    chunks, i = [], 0
    while i < len(text):
        chunks.append(text[i:i + max_chars])
        i += max_chars - overlap
    return chunks


def extract_api_endpoints(raw_text: str) -> str | None:
    if not raw_text or len(raw_text) < 10:
        return None
    
    import re
    found = set()
    lines = raw_text.split('\n')
    
    # Pattern 1: explicit METHOD /path
    explicit_regex = re.compile(r'\b(GET|POST|PUT|PATCH|DELETE)\s+([^\s"\'\>\|]*\/[^\s"\'\>\|]*)', re.IGNORECASE)
    
    # Pattern 3: YAML/OpenAPI style
    yaml_path_regex = re.compile(r'^(\s{0,4})(\/[A-Za-z0-9\-._~:/?#\[\]@!$&\'()*+,;=%{}]{2,}):\s*$')
    yaml_method_regex = re.compile(r'^\s{2,8}(get|post|put|patch|delete):\s*$', re.IGNORECASE)
    
    pending_yaml_path = None
    
    def has_unbalanced_curly_braces(path: str) -> bool:
        has_open = '{' in path
        has_close = '}' in path
        if has_open != has_close:
            return True
        if not has_open:
            return False
        return bool(re.search(r'\{[^}]*$', path) or re.search(r'^[^{]*\}', path))

    for line in lines:
        trimmed = line.strip()
        
        # YAML path match
        yaml_match = yaml_path_regex.match(line)
        if yaml_match:
            pending_yaml_path = yaml_match.group(2)
            continue
            
        # YAML method match
        if pending_yaml_path:
            yaml_method_match = yaml_method_regex.match(line)
            if yaml_method_match:
                method = yaml_method_match.group(1).upper()
                if not has_unbalanced_curly_braces(pending_yaml_path):
                    found.add(f"{method} {pending_yaml_path}")
            elif trimmed and not trimmed.startswith('#'):
                pending_yaml_path = None
                
        # Regex search for explicit ones
        for match in explicit_regex.finditer(line):
            method = match.group(1).upper()
            path = match.group(2)
            # clean up path trailing characters like braces or commas or quotes
            path = re.sub(r'[\.,;:\"\'\)>\]]+$', '', path)
            if not has_unbalanced_curly_braces(path):
                found.add(f"{method} {path}")
                
    if found:
        return "\n".join(sorted(found))
    return None
